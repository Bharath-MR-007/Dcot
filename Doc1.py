from docx import Document
document = Document()

document.add_heading('Riyadh VPN Change', 0)

document.add_heading('Pre Implementation Plan', level=1)

p = document.add_paragraph('Login to Riyadh APIC GUI, x:   ')

url=p.add_run('https://10.200.4.15')
url.bold = True
url.italic = True
print(x)
#p.add_run(' Riyadh APIC ')

document.add_paragraph(
    'The health of controller should be fully fit:  Go to System -> controllers (take snapshot as below)', style='List Number'
)
document.add_paragraph(
    'Go to System -> Faults (screenshot)  & Go to System -> Dashboard (Take snapshot as below)', style='List Number'
)
document.add_paragraph(
    'Take backups of the APIC fabric. (Riyadh- https://10.200.4.15) \n Navigate to Admin -> Config Rollbacks  Select tenants as “All fabric” \n Put CHG# in the Description Field >> Click on "Create A snapshot Now"', style='List Number'
)



document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)



table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'


document.add_page_break()

document.save('vpn-Change.docx')

